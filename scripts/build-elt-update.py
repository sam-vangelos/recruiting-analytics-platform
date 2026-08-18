"""Render the weekly ELT Recruiting Update from elt-data.json.

Reproduces the legacy ELT doc format exactly (Hires block + three Role Progress
sections with conducted/passed funnels and QTD offers), but from live data — so
it can't inherit the manual doc's copy-paste errors (wrong stream splits,
mismatched date ranges, skipped stage numbers, inconsistent labels).

Outputs a .docx (the real medium — paste into the ELT doc) and an .html
preview. Business logic lives in the TS generator; this only formats.

Usage: python3 scripts/build-elt-update.py <elt-data.json> <out.docx> <out.html>
"""
import json
import sys
import html

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INK = RGBColor(0x1F, 0x29, 0x33)
MUTED = RGBColor(0x7B, 0x84, 0x94)
SLATE = RGBColor(0x24, 0x34, 0x4D)

WORDS = {0: "no", 1: "one", 2: "two", 3: "three", 4: "four", 5: "five", 6: "six",
         7: "seven", 8: "eight", 9: "nine", 10: "ten"}


def split_str(subs, key="count"):
    return "(" + ", ".join(f"{s['label']} {s[key]}" for s in subs) + ")"


def stage_line(idx, stage):
    """'2. Manager/Tech Screen Conducted - 4 (PE 0, FDE 4): 3 passed'"""
    prefix = f"{idx}. " if idx else ""
    line = f"{prefix}{stage['label']} Conducted - {stage['conducted']} {split_str(stage['subs'], 'conducted')}"
    if stage["conducted"]:
        line += f": {stage['passed']} passed"
    return line


# ------------------------------------------------------------------- docx
def hire_line(h):
    """Legacy format: Role - Location - Candidate - Department (Priority) - Start date."""
    parts = [h["role"]]
    if h.get("location"):
        parts.append(h["location"])
    parts.append(h["candidate"])
    dept = h.get("department")
    prio = h.get("priority")
    if dept and prio:
        parts.append(f"{dept} ({prio})")
    elif dept:
        parts.append(dept)
    elif prio:
        parts.append(f"({prio})")
    parts.append(f"Start date {h['startsOn'] or 'TBD'}")
    return " - ".join(parts) + "."


def build_docx(data, out):
    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def head(text, size=15, color=SLATE, space_before=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = color
        return p

    def line(text, bold=False, color=INK, indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if indent:
            p.paragraph_format.left_indent = Pt(indent)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(11)
        r.font.color.rgb = color
        return p

    head(data["weekLabel"], size=16, color=INK, space_before=0)

    # Hires
    head(f"Hires: (Offer Accepted b/w {data['weekShort']})", size=13)
    n = len(data["hires"])
    line(f"We had {WORDS.get(n, n)} offer{'' if n == 1 else 's'} accepted this week.", bold=True)
    for h in data["hires"]:
        line(hire_line(h))
    if not data["hires"]:
        line("No offers accepted this week.", color=MUTED)
    line(data["hiresNote"], color=MUTED)

    # Role Progress sections
    for section in data["sections"]:
        head(f"{section['title']} Role Progress b/w {data['weekShort']}", size=12, color=INK)
        q = section["qtdOffers"]
        names = f" - {', '.join(q['names'])}" if q["names"] else ""
        line(f"QTD Offer Accepted - {q['total']} {split_str(q['subs'])}{names}")
        for i, stage in enumerate(section["stages"]):
            # RPS unnumbered (as in the legacy doc), then 1..4
            line(stage_line(i if i else 0, stage) if i else stage_line(0, stage))
        w = section["weekOffers"]
        wnames = f" - {', '.join(w['names'])}" if w["names"] else ""
        line(f"5. Offer Accepted - {w['total']} {split_str(w['subs'])}{wnames}")
    doc.save(out)


# ------------------------------------------------------------------- html
def build_html(data, out):
    e = lambda v: html.escape(str(v))
    css = """<style>
    body{font-family:Calibri,'Segoe UI',-apple-system,sans-serif;color:#1F2933;margin:0;background:#fff}
    .wrap{max-width:820px;margin:0 auto;padding:30px 36px 48px}
    .banner{background:#FFF8E6;border:1px solid #EAD79B;border-radius:6px;padding:8px 12px;font-size:11.5px;color:#6b5900;margin-bottom:20px}
    h1{font-size:20px;margin:0 0 14px;color:#1F2933}
    h2{font-size:14px;color:#24344D;margin:20px 0 4px}
    .lead{font-weight:700;margin:2px 0 6px}
    .hire{margin:1px 0;font-size:13px}
    .sec{font-size:13.5px;font-weight:700;color:#1F2933;margin:18px 0 4px;padding-top:10px;border-top:1px solid #E4E7EB}
    .row{margin:1px 0;font-size:12.5px}
    .qtd{color:#2E7D46}
    .mut{color:#7B8794;font-size:11.5px}
    .stagenum{color:#7B8794}
    .pass{color:#2E7D46}
    </style>"""

    def split_html(subs, key="count"):
        return "(" + ", ".join(f"{e(s['label'])} {s[key]}" for s in subs) + ")"

    parts = [css, '<div class="wrap">']
    parts.append('<div class="banner"><b>PREVIEW</b> — live-data ELT update in the legacy format, generated '
                 + e(data["generatedAt"][:16].replace("T", "  ")) + ' UTC. Draft to review and paste.</div>')
    parts.append(f'<h1>{e(data["weekLabel"])}</h1>')
    parts.append(f'<h2>Hires: (Offer Accepted b/w {e(data["weekShort"])})</h2>')
    n = len(data["hires"])
    parts.append(f'<div class="lead">We had {WORDS.get(n, n)} offer{"" if n == 1 else "s"} accepted this week.</div>')
    for h in data["hires"]:
        parts.append(f'<div class="hire">{e(hire_line(h))}</div>')
    if not data["hires"]:
        parts.append('<div class="hire mut">No offers accepted this week.</div>')
    parts.append(f'<div class="mut" style="margin-top:6px">{e(data["hiresNote"])}</div>')

    for section in data["sections"]:
        parts.append(f'<div class="sec">{e(section["title"])} Role Progress b/w {e(data["weekShort"])}</div>')
        q = section["qtdOffers"]
        names = f' — {e(", ".join(q["names"]))}' if q["names"] else ""
        parts.append(f'<div class="row qtd">QTD Offer Accepted — {q["total"]} {split_html(q["subs"])}{names}</div>')
        for i, stage in enumerate(section["stages"]):
            num = f'<span class="stagenum">{i}.</span> ' if i else ""
            passed = f': <span class="pass">{stage["passed"]} passed</span>' if stage["conducted"] else ""
            parts.append(f'<div class="row">{num}{e(stage["label"])} Conducted — {stage["conducted"]} {split_html(stage["subs"], "conducted")}{passed}</div>')
        w = section["weekOffers"]
        wnames = f' — {e(", ".join(w["names"]))}' if w["names"] else ""
        parts.append(f'<div class="row"><span class="stagenum">5.</span> Offer Accepted — {w["total"]} {split_html(w["subs"])}{wnames}</div>')
    parts.append("</div>")
    open(out, "w").write("".join(parts))


def main():
    src, docx_out, html_out = sys.argv[1], sys.argv[2], sys.argv[3]
    with open(src) as fh:
        data = json.load(fh)
    build_docx(data, docx_out)
    build_html(data, html_out)
    n_off = sum(s["weekOffers"]["total"] for s in data["sections"])
    print(f"wrote {docx_out} + {html_out} — week {data['weekShort']}, {len(data['hires'])} hires, {n_off} focus-req offers")


if __name__ == "__main__":
    main()
